import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

def create_docx_from_md(md_content: str, docx_path: str):
    """
    Convert Markdown content to DOCX, supporting images.
    """
    doc = Document()
    
    # 设置中文字体
    try:
        style = doc.styles['Normal']
        style.font.name = u'微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    except Exception as e:
        logger.warning(f"Failed to set font: {e}")
    
    lines = md_content.splitlines()
    
    table_mode = False
    table_data = []
    
    # Regex for images: ![alt](path)
    img_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    
    def flush_table():
        nonlocal table_mode, table_data
        if table_data:
            create_table(doc, table_data)
        table_mode = False
        table_data = []

    for line in lines:
        line = line.strip()
        
        # 处理空行
        if not line:
            if table_mode:
                flush_table()
            continue
            
        # 标题处理
        if line.startswith('# '):
            if table_mode: flush_table()
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            if table_mode: flush_table()
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            if table_mode: flush_table()
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            if table_mode: flush_table()
            doc.add_heading(line[5:], level=3)
            
        # 列表处理
        elif line.startswith('* ') or line.startswith('- '):
            if table_mode: flush_table()
            # Handle basic list item
            p = doc.add_paragraph(line[2:], style='List Bullet')
            # Handle bold inside list
            # Note: Complex formatting inside list items is tricky with python-docx simple adding
            # But we can try to apply bold formatting if detected
            
        # 表格处理
        elif line.startswith('|'):
            if not table_mode:
                table_mode = True
            # Strip outer pipes if they exist
            row_content = line.strip().strip('|').split('|')
            if '---' in line:
                continue
            table_data.append([cell.strip() for cell in row_content])
            
        # 图片处理 (Assume image is on its own line for now as per plugin logic)
        elif line.startswith('![') and '](' in line and line.endswith(')'):
            if table_mode: flush_table()
            match = img_pattern.match(line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                
                # Clean path
                img_path = img_path.strip()
                
                # Check if path exists
                if os.path.exists(img_path):
                    try:
                        # Add picture
                        # Use a width that fits page (e.g. 6 inches)
                        doc.add_picture(img_path, width=Inches(6.0))
                        
                        # Add caption if meaningful
                        if alt_text and alt_text.lower() != "mermaid diagram":
                             caption = doc.add_paragraph(alt_text)
                             caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                             
                        logger.info(f"Inserted image: {img_path}")
                    except Exception as e:
                        logger.error(f"Failed to add picture {img_path}: {e}")
                        doc.add_paragraph(f"[Image: {alt_text}]")
                else:
                    logger.warning(f"Image not found: {img_path}")
                    doc.add_paragraph(f"[Image not found: {alt_text}]")
            else:
                # Regex didn't match fully but looked like image, treat as text
                doc.add_paragraph(line)

        # 普通段落
        else:
            if table_mode: flush_table()
            
            p = doc.add_paragraph()
            # Handle bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    if table_mode:
        flush_table()

    doc.save(docx_path)
    return True

def create_table(doc, data):
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0])
    # Ensure cols matches max row length? docx add_table requires fixed cols
    # But usually markdown tables are consistent.
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell_text

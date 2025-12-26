import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

def create_docx_from_md(md_path, docx_path):
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    table_mode = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        
        # 处理空行
        if not line:
            if table_mode:
                # 表格结束，渲染表格
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            continue
            
        # 标题处理
        if line.startswith('# '):
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            doc.add_heading(line[4:], level=2)
            
        # 列表处理
        elif line.startswith('* ') or line.startswith('- '):
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            doc.add_paragraph(line[2:], style='List Bullet')
            
        # 表格处理
        elif line.startswith('|'):
            if not table_mode:
                table_mode = True
            # 去掉首尾的 |
            row_content = line.strip('|').split('|')
            # 过滤掉分隔行 (e.g., :---)
            if '---' in line:
                continue
            table_data.append([cell.strip() for cell in row_content])
            
        # 普通段落
        else:
            if table_mode:
                create_table(doc, table_data)
                table_mode = False
                table_data = []
            # 简单处理粗体 **text**
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # 如果文件结束时还在表格模式
    if table_mode:
        create_table(doc, table_data)

    try:
        doc.save(docx_path)
        # print(f"Successfully created {docx_path}")
        return True, f"Successfully created {docx_path}"
    except Exception as e:
        sys.stderr.write(f"Error saving docx: {e}\n")
        return False, str(e)

def create_table(doc, data):
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            # 处理越界情况
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell_text

if __name__ == '__main__':
    md_file = r"e:\workspace_seerlord\test\附件下载审批功能实施方案.md"
    docx_file = r"e:\workspace_seerlord\test\attachment_approval_plan.docx"
    create_docx_from_md(md_file, docx_file)
